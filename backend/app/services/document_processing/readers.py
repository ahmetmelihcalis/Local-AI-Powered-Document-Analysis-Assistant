from dataclasses import dataclass
from io import BytesIO
from pathlib import Path

import pymupdf
from docx import Document


TEXT_EXTENSIONS = {".txt", ".md"}


@dataclass
class PageText:
    page: int
    text: str


@dataclass
class DocumentText:
    text: str
    section: str | None = None


def read_text_document(filename: str, content: bytes) -> str:
    extension = Path(filename).suffix.lower()

    if extension not in TEXT_EXTENSIONS:
        raise ValueError("Only TXT and Markdown files are supported at this stage.")

    for encoding in ("utf-8-sig", "utf-8", "cp1254"):
        try:
            text = content.decode(encoding)
            break
        except UnicodeDecodeError:
            continue
    else:
        raise ValueError("The file encoding could not be read.")

    text = text.replace("\x00", "").strip()

    if not text:
        raise ValueError("The document does not contain readable text.")

    return text


def read_pdf_document(content: bytes) -> list[PageText]:
    try:
        document = pymupdf.open(stream=content, filetype="pdf")
    except Exception as error:
        raise ValueError("The PDF file could not be read.") from error

    if document.needs_pass:
        document.close()
        raise ValueError("Encrypted PDF files are not supported.")

    try:
        pages = [
            PageText(page=number, text=page.get_text("text", sort=True).strip())
            for number, page in enumerate(document, start=1)
        ]
    finally:
        document.close()

    pages = [page for page in pages if page.text]

    if not pages:
        raise ValueError("The PDF does not contain a readable text layer.")

    return pages


def read_docx_document(content: bytes) -> list[DocumentText]:
    try:
        document = Document(BytesIO(content))
    except Exception as error:
        raise ValueError("The DOCX file could not be read.") from error

    blocks: list[DocumentText] = []
    current_section: str | None = None

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()

        if not text:
            continue

        if paragraph.style and paragraph.style.name.startswith("Heading"):
            current_section = text

        blocks.append(DocumentText(text=text, section=current_section))

    for table in document.tables:
        for row in table.rows:
            text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if text:
                blocks.append(DocumentText(text=text, section=current_section))

    if not blocks:
        raise ValueError("The DOCX file does not contain readable text.")

    return blocks
